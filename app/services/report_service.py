from datetime import datetime
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_excel_report(
    report_folder: Path,
    filename: str,
    anova_result: dict,
    tukey_results: list[dict],
    selected_discipline: str,
    metrics: dict,
) -> Path:
    """
    Создаёт Excel-отчёт по результатам ANOVA и Tukey HSD.
    """

    report_folder.mkdir(parents=True, exist_ok=True)

    safe_filename = _build_report_filename(filename)
    report_path = report_folder / safe_filename

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "ANOVA-отчёт"

    _setup_page(worksheet)
    _write_title(worksheet)
    _write_analysis_params(
        worksheet=worksheet,
        anova_result=anova_result,
        selected_discipline=selected_discipline,
        metrics=metrics,
    )
    _write_anova_table(worksheet, anova_result)
    _write_groups_summary(worksheet, anova_result)
    _write_tukey_table(worksheet, tukey_results)
    _write_conclusion(worksheet, anova_result)

    workbook.save(report_path)

    return report_path


def create_docx_report(
    report_folder: Path,
    filename: str,
    anova_result: dict,
    tukey_results: list[dict],
    selected_discipline: str,
    metrics: dict,
) -> Path:
    """
    Создаёт Word-отчёт по результатам ANOVA и Tukey HSD.
    """

    report_folder.mkdir(parents=True, exist_ok=True)

    safe_filename = _build_report_filename_docx(filename)
    report_path = report_folder / safe_filename

    document = Document()

    # Заголовок
    title = document.add_heading(
        "Отчёт по ANOVA-анализу успеваемости студентов", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph(
        "Однофакторный дисперсионный анализ и post-hoc тест Tukey HSD"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(12)

    document.add_paragraph()

    # Параметры анализа
    document.add_heading("Параметры анализа", level=1)

    params = [
        ("Результативный показатель", anova_result["score_column"]),
        ("Фактор", anova_result["factor_column"]),
        ("Дисциплина", selected_discipline if selected_discipline else "Все дисциплины"),
        ("Уровень значимости", anova_result["alpha"]),
        ("Количество записей", metrics["records_count"]),
        ("Количество групп", metrics["groups_count"]),
        ("Количество методов обучения", metrics["methods_count"]),
        ("Средний балл", metrics["average_score"]),
    ]

    for name, value in params:
        p = document.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(str(value))

    document.add_paragraph()

    # Таблица ANOVA
    document.add_heading("Таблица ANOVA", level=1)

    anova_table = document.add_table(rows=1, cols=6)
    anova_table.style = "Light Grid Accent 1"

    headers = anova_table.rows[0].cells
    header_texts = ["Источник вариации", "SS", "df", "MS", "F", "p-value"]

    for i, text in enumerate(header_texts):
        headers[i].text = text
        _set_cell_background(headers[i], "DCE9FB")

    for item in anova_result["anova_table"]:
        row_cells = anova_table.add_row().cells
        row_cells[0].text = item["source"]
        row_cells[1].text = str(item["ss"])
        row_cells[2].text = str(item["df"])
        row_cells[3].text = str(item["ms"])
        row_cells[4].text = str(item["f"])
        row_cells[5].text = str(item["p_value"])

    document.add_paragraph()

    # Средние значения по группам фактора
    document.add_heading("Средние значения по группам фактора", level=1)

    groups_table = document.add_table(rows=1, cols=4)
    groups_table.style = "Light Grid Accent 1"

    headers = groups_table.rows[0].cells
    header_texts = ["Группа", "Количество", "Среднее", "Стандартное отклонение"]

    for i, text in enumerate(header_texts):
        headers[i].text = text
        _set_cell_background(headers[i], "DCE9FB")

    for group in anova_result["groups_summary"]:
        row_cells = groups_table.add_row().cells
        row_cells[0].text = group["name"]
        row_cells[1].text = str(group["count"])
        row_cells[2].text = str(group["mean"])
        row_cells[3].text = str(group["std"])

    document.add_paragraph()

    # Post-hoc тест Tukey HSD
    document.add_heading("Post-hoc тест Tukey HSD", level=1)

    tukey_table = document.add_table(rows=1, cols=7)
    tukey_table.style = "Light Grid Accent 1"

    headers = tukey_table.rows[0].cells
    header_texts = [
        "Группа 1",
        "Группа 2",
        "Разница средних",
        "p-value",
        "Нижняя граница",
        "Верхняя граница",
        "Значимо",
    ]

    for i, text in enumerate(header_texts):
        headers[i].text = text
        _set_cell_background(headers[i], "DCE9FB")

    for item in tukey_results:
        row_cells = tukey_table.add_row().cells
        row_cells[0].text = item["group_1"]
        row_cells[1].text = item["group_2"]
        row_cells[2].text = str(item["mean_difference"])
        row_cells[3].text = str(item["p_value"])
        row_cells[4].text = str(item["lower_bound"])
        row_cells[5].text = str(item["upper_bound"])
        row_cells[6].text = item["significant_text"]

    document.add_paragraph()

    # Итоговый вывод
    document.add_heading("Итоговый вывод", level=1)

    conclusion_p = document.add_paragraph(anova_result["conclusion"])

    if anova_result["significant"]:
        result_text = (
            f"Итог: выбранный фактор «{anova_result['factor_column']}» "
            f"статистически значимо влияет на показатель «{anova_result['score_column']}»."
        )
    else:
        result_text = (
            f"Итог: статистически значимого влияния фактора "
            f"«{anova_result['factor_column']}» на показатель "
            f"«{anova_result['score_column']}» не выявлено."
        )

    result_p = document.add_paragraph(result_text)
    result_p.runs[0].bold = True

    document.save(report_path)

    return report_path


def _build_report_filename(source_filename: str) -> str:
    """
    Формирует имя файла отчёта для Excel.
    """

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    clean_name = Path(source_filename).stem
    clean_name = "".join(
        char if char.isalnum() or char in ("_", "-") else "_"
        for char in clean_name
    )

    return f"anova_report_{clean_name}_{timestamp}.xlsx"


def _build_report_filename_docx(source_filename: str) -> str:
    """
    Формирует имя файла отчёта для Word.
    """

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    clean_name = Path(source_filename).stem
    clean_name = "".join(
        char if char.isalnum() or char in ("_", "-") else "_"
        for char in clean_name
    )

    return f"anova_report_{clean_name}_{timestamp}.docx"


def _setup_page(worksheet) -> None:
    """
    Настраивает ширину столбцов и базовое оформление листа.
    """

    column_widths = {
        "A": 28,
        "B": 24,
        "C": 18,
        "D": 20,
        "E": 18,
        "F": 18,
        "G": 18,
    }

    for column, width in column_widths.items():
        worksheet.column_dimensions[column].width = width

    worksheet.freeze_panes = "A5"


def _write_title(worksheet) -> None:
    """
    Записывает заголовок отчёта.
    """

    worksheet.merge_cells("A1:G1")
    title_cell = worksheet["A1"]
    title_cell.value = "Отчёт по ANOVA-анализу успеваемости студентов"
    title_cell.font = Font(bold=True, size=16)
    title_cell.alignment = Alignment(horizontal="center")

    worksheet.merge_cells("A2:G2")
    subtitle_cell = worksheet["A2"]
    subtitle_cell.value = "Однофакторный дисперсионный анализ и post-hoc тест Tukey HSD"
    subtitle_cell.font = Font(size=12)
    subtitle_cell.alignment = Alignment(horizontal="center")


def _write_analysis_params(
    worksheet,
    anova_result: dict,
    selected_discipline: str,
    metrics: dict,
) -> None:
    """
    Записывает параметры анализа.
    """

    start_row = 4

    worksheet[f"A{start_row}"] = "Параметры анализа"
    worksheet[f"A{start_row}"].font = Font(bold=True, size=13)

    params = [
        ("Результативный показатель", anova_result["score_column"]),
        ("Фактор", anova_result["factor_column"]),
        ("Дисциплина", selected_discipline if selected_discipline else "Все дисциплины"),
        ("Уровень значимости", anova_result["alpha"]),
        ("Количество записей", metrics["records_count"]),
        ("Количество групп", metrics["groups_count"]),
        ("Количество методов обучения", metrics["methods_count"]),
        ("Средний балл", metrics["average_score"]),
    ]

    row = start_row + 1

    for name, value in params:
        worksheet[f"A{row}"] = name
        worksheet[f"B{row}"] = value
        worksheet[f"A{row}"].font = Font(bold=True)
        row += 1


def _write_anova_table(worksheet, anova_result: dict) -> None:
    """
    Записывает таблицу ANOVA.
    """

    start_row = 15

    worksheet[f"A{start_row}"] = "Таблица ANOVA"
    worksheet[f"A{start_row}"].font = Font(bold=True, size=13)

    headers = ["Источник вариации", "SS", "df", "MS", "F", "p-value"]
    _write_table_header(worksheet, start_row + 1, headers)

    row = start_row + 2

    for item in anova_result["anova_table"]:
        values = [
            item["source"],
            item["ss"],
            item["df"],
            item["ms"],
            item["f"],
            item["p_value"],
        ]

        _write_table_row(worksheet, row, values)
        row += 1


def _write_groups_summary(worksheet, anova_result: dict) -> None:
    """
    Записывает краткую статистику по группам фактора.
    """

    start_row = 22

    worksheet[f"A{start_row}"] = "Средние значения по группам фактора"
    worksheet[f"A{start_row}"].font = Font(bold=True, size=13)

    headers = ["Группа", "Количество", "Среднее", "Стандартное отклонение"]
    _write_table_header(worksheet, start_row + 1, headers)

    row = start_row + 2

    for group in anova_result["groups_summary"]:
        values = [
            group["name"],
            group["count"],
            group["mean"],
            group["std"],
        ]

        _write_table_row(worksheet, row, values)
        row += 1


def _write_tukey_table(worksheet, tukey_results: list[dict]) -> None:
    """
    Записывает таблицу Tukey HSD.
    """

    start_row = 32

    worksheet[f"A{start_row}"] = "Post-hoc тест Tukey HSD"
    worksheet[f"A{start_row}"].font = Font(bold=True, size=13)

    headers = [
        "Группа 1",
        "Группа 2",
        "Разница средних",
        "p-value",
        "Нижняя граница",
        "Верхняя граница",
        "Значимо",
    ]

    _write_table_header(worksheet, start_row + 1, headers)

    row = start_row + 2

    for item in tukey_results:
        values = [
            item["group_1"],
            item["group_2"],
            item["mean_difference"],
            item["p_value"],
            item["lower_bound"],
            item["upper_bound"],
            item["significant_text"],
        ]

        _write_table_row(worksheet, row, values)
        row += 1


def _write_conclusion(worksheet, anova_result: dict) -> None:
    """
    Записывает итоговый вывод.
    """

    start_row = 45

    worksheet[f"A{start_row}"] = "Итоговый вывод"
    worksheet[f"A{start_row}"].font = Font(bold=True, size=13)

    worksheet.merge_cells(start_row=start_row + 1, start_column=1, end_row=start_row + 3, end_column=7)

    conclusion_cell = worksheet[f"A{start_row + 1}"]
    conclusion_cell.value = anova_result["conclusion"]
    conclusion_cell.alignment = Alignment(wrap_text=True, vertical="top")

    if anova_result["significant"]:
        result_text = (
            f"Итог: выбранный фактор «{anova_result['factor_column']}» "
            f"статистически значимо влияет на показатель «{anova_result['score_column']}»."
        )
    else:
        result_text = (
            f"Итог: статистически значимого влияния фактора "
            f"«{anova_result['factor_column']}» на показатель "
            f"«{anova_result['score_column']}» не выявлено."
        )

    worksheet.merge_cells(start_row=start_row + 5, start_column=1, end_row=start_row + 6, end_column=7)

    result_cell = worksheet[f"A{start_row + 5}"]
    result_cell.value = result_text
    result_cell.font = Font(bold=True)
    result_cell.alignment = Alignment(wrap_text=True, vertical="top")


def _write_table_header(worksheet, row: int, headers: list[str]) -> None:
    """
    Записывает строку заголовков таблицы.
    """

    header_fill = PatternFill(
        start_color="DCE9FB",
        end_color="DCE9FB",
        fill_type="solid"
    )

    for column_index, header in enumerate(headers, start=1):
        cell = worksheet.cell(row=row, column=column_index)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = _thin_border()


def _write_table_row(worksheet, row: int, values: list) -> None:
    """
    Записывает обычную строку таблицы.
    """

    for column_index, value in enumerate(values, start=1):
        cell = worksheet.cell(row=row, column=column_index)
        cell.value = value
        cell.alignment = Alignment(vertical="center", wrap_text=True)
        cell.border = _thin_border()


def _set_cell_background(cell, color: str) -> None:
    """
    Устанавливает цвет фона для ячейки в Word документе.
    """

    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def _thin_border() -> Border:
    """
    Возвращает тонкую границу для ячеек таблиц.
    """

    side = Side(style="thin", color="B7B7B7")

    return Border(
        left=side,
        right=side,
        top=side,
        bottom=side
    )